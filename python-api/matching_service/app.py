
import os, re, json, time, logging, unicodedata, tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import requests as http_requests    # renommé pour éviter conflits avec FastAPI
import mysql.connector
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from groq import Groq, RateLimitError
from dotenv import load_dotenv

load_dotenv()

# ================================================================
#  CONFIG
# ================================================================

DB = {
    "host":     os.getenv("DB_HOST",     "localhost"),
    "database": os.getenv("DB_NAME",     "projetb"),
    "user":     os.getenv("DB_USER",     "root"),
    "password": os.getenv("DB_PASSWORD", ""),
    "charset":  "utf8mb4",
}

GROQ_API_KEY  = os.getenv("GROQ_API_KEY", "")
RACINE_PROJET = os.getenv("RACINE_PROJET", "")   # inutile sur Render, conservé legacy
TESSERACT_CMD = os.getenv("TESSERACT_CMD", "")
MODELE        = "llama-3.3-70b-versatile"
SCORE_MIN     = 10

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("matching")

try:
    _t = mysql.connector.connect(**DB); _t.close()
    log.info(f"MySQL OK — base : {DB['database']}")
except Exception as e:
    log.error(f"MySQL ERREUR : {e}")

groq_client = Groq(api_key=GROQ_API_KEY)

app = FastAPI(title="SmartRecruit Matching IA", version="7.2", docs_url="/docs")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"], allow_credentials=True)


# ================================================================
#  MODÈLES PYDANTIC
# ================================================================

class RequeteChat(BaseModel):
    question:        str  = Field(...)
    recruteur_id:    Optional[int] = None
    historique_chat: list = Field(default=[])

class RequeteExtractionCV(BaseModel):
    texte_cv:    str
    candidat_id: int

class RequeteAnalyserFichier(BaseModel):
    candidat_id: int
    chemin_cv:   str   # accepte chemin local OU URL Cloudinary

class RequeteUploadCV(BaseModel):
    candidat_id:    int
    chemin_cv:      str = Field(default="")   # chemin local (legacy)
    cloudinary_url: str = Field(default="")   # URL Cloudinary (v7.2)
    texte_cv:       str = Field(default="")   # texte pré-extrait (optionnel)


# ================================================================
#  GROQ — APPEL CENTRALISÉ AVEC RETRY 429
# ================================================================

def groq_appel(messages: list, max_tokens: int = 300,
               temperature: float = 0.2, json_mode: bool = False) -> str:
    kwargs = dict(model=MODELE, messages=messages,
                  max_tokens=max_tokens, temperature=temperature)
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    for tentative in range(1, 4):
        try:
            resp = groq_client.chat.completions.create(**kwargs)
            return resp.choices[0].message.content or ""
        except RateLimitError as e:
            m = re.search(r"try again in (\d+)m([\d.]+)s", str(e))
            attente = min((int(m.group(1))*60 + float(m.group(2))) if m else 30, 60)
            log.warning(f"Rate limit — tentative {tentative}/3, attente {attente:.0f}s")
            if tentative == 3:
                raise HTTPException(429, "Limite Groq atteinte. Réessayez dans quelques minutes.")
            time.sleep(attente)
        except Exception as e:
            raise HTTPException(500, f"Erreur Groq : {e}")
    return ""

def nettoyer_json(texte: str) -> str:
    texte = re.sub(r"^```(?:json)?\s*", "", texte, flags=re.MULTILINE)
    texte = re.sub(r"\s*```$",          "", texte, flags=re.MULTILINE)
    texte = texte.strip()
    debut = texte.find("{"); fin = texte.rfind("}")
    if debut != -1 and fin > debut:
        return texte[debut:fin+1]
    return texte


# ================================================================
#  UTILITAIRE — SUPPRESSION DES ACCENTS (v7.1)
# ================================================================

def sans_accent(texte: str) -> str:
    """Retire les accents : 'Doué' → 'Doue', 'élève' → 'eleve'."""
    return ''.join(
        c for c in unicodedata.normalize('NFD', texte)
        if unicodedata.category(c) != 'Mn'
    )


# ================================================================
#  CLASSIFICATION PYTHON PUR — 0 TOKEN GROQ
# ================================================================

CORRECTIONS = {
    r"\bchaiche\b":       "cherche",
    r"\bchèrche\b":       "cherche",
    r"\bcomaercial\b":    "commercial",
    r"\bcommercail\b":    "commercial",
    r"\binformaticein\b": "informaticien",
    r"\byamossoukro\b":   "yamoussoukro",
    r"\byamoussokro\b":   "yamoussoukro",
    r"\babidgane?\b":     "abidjan",
    r"\bdeveloppeure?\b": "développeur",
    r"\belectriciene?\b": "électricien",
    r"\bsecretaire\b":    "secrétaire",
    r"\bcomptabilite\b":  "comptabilité",
}

MOTS_CONVERSATION = {
    "bonjour","bonsoir","salut","hello","hi","allo",
    "aide","help","comment","fonctionner","fonctionne",
    "merci","svp","stp","qu'est-ce",
}

MOTS_STATISTIQUE = {
    "combien","nombre","total","statistique","stat",
    "résumé","liste complète","tous les candidats","répartition","moyenne",
}

STOP_WORDS = {
    "je","un","une","des","les","le","la","de","du","et","en",
    "que","qui","avec","pour","dans","sur","cherche","recherche",
    "voudrais","veux","besoin","trouver","profil","candidat","cv",
    "voir","avoir","plus","moins","ans","année","années","mois",
    "aime","aimer","avoir","être","faire","savoir","pouvoir",
}

VILLES_CONNUES = {
    "abidjan","yamoussoukro","bouaké","bouake","daloa",
    "san pedro","korhogo","man","divo","gagnoa","abengourou",
    "bondoukou","odienné","dakar","accra","lomé","cotonou",
    "bamako","paris","lyon","marseille","bordeaux",
}

MOTS_FEMME = {"femme","féminin","féminine","dame","madame","mlle"}
MOTS_HOMME = {"homme","masculin","monsieur"}

METIERS = [
    "développeur","developer","dev","ingénieur","engineer",
    "comptable","commercial","médecin","infirmier","infirmière",
    "mécanicien","plombier","électricien","secrétaire","manager",
    "directeur","directrice","technicien","technicienne",
    "informaticien","data scientist","data analyst","analyste",
    "consultant","graphiste","designer","architecte","juriste",
    "avocat","chargé de communication","chargé de marketing",
    "responsable","chef de projet","full stack","fullstack",
    "frontend","backend","devops","sysadmin","pharmacien",
    "enseignant","professeur","formateur","logisticien",
    "community manager","rh","ressources humaines","recruteur",
]

TECHS = [
    "php","mysql","javascript","python","java","react","vue",
    "angular","nodejs","laravel","symfony","django","flask",
    "html","css","sql","postgresql","mongodb","redis",
    "docker","kubernetes","aws","azure","git","linux",
    "word","excel","powerpoint","powerbi","tableau",
    "photoshop","illustrator","figma","sap","sage",
    "c++","c#",".net","kotlin","swift","flutter",
    "machine learning","deep learning","tensorflow",
]


def normaliser(question: str) -> str:
    q = question.lower().strip()
    for pattern, rep in CORRECTIONS.items():
        q = re.sub(pattern, rep, q, flags=re.IGNORECASE)
    return q


def classifier(question_norm: str) -> str:
    mots = set(re.findall(r'\b\w+\b', question_norm.lower()))
    if mots & MOTS_STATISTIQUE:
        return "statistique"
    if mots & MOTS_CONVERSATION:
        has_metier = any(
            re.search(r'\b' + re.escape(m) + r'\b', question_norm, re.IGNORECASE)
            for m in METIERS[:20]
        )
        if not has_metier:
            return "conversation"
    return "recherche"


def extraire_criteres(question_norm: str, question_orig: str) -> dict:
    criteres = {
        "genre": None, "ville": None, "exp_min": None, "exp_max": None,
        "poste": None, "competences": [], "mots_cles_libres": [],
        "disponibilite": None, "niveau": None, "prenom_nom": None,
    }
    q = question_norm.lower()

    mots_q = set(re.findall(r'\b\w+\b', q))
    if mots_q & MOTS_FEMME: criteres["genre"] = "femme"
    elif mots_q & MOTS_HOMME: criteres["genre"] = "homme"

    m = re.search(r'(\d+)\s*(?:à|-)\s*(\d+)\s*ans', q)
    if m:
        criteres["exp_min"] = int(m.group(1)); criteres["exp_max"] = int(m.group(2))
    else:
        m = re.search(r'(\d+)\s*ans', q)
        if m:
            val = int(m.group(1))
            criteres["exp_min"] = val if re.search(r'plus\s*de|minimum|au\s*moins', q) else max(0, val-1)

    if re.search(r'imm[eé]diat|disponible\s*maintenant', q):
        criteres["disponibilite"] = "immediat"
    elif re.search(r'1\s*mois|dans\s*un\s*mois', q):
        criteres["disponibilite"] = "1_mois"
    elif re.search(r'3\s*mois', q):
        criteres["disponibilite"] = "3_mois"

    if re.search(r'master|bac\+5', q):    criteres["niveau"] = "Master (Bac+5)"
    elif re.search(r'licence|bac\+3', q): criteres["niveau"] = "Licence (Bac+3)"
    elif re.search(r'ing[eé]nieur', q):   criteres["niveau"] = "Ingénieur"
    elif re.search(r'bts|bac\+2', q):     criteres["niveau"] = "BTS"

    for ville in VILLES_CONNUES:
        if re.search(r'\b' + re.escape(ville) + r'\b', q):
            criteres["ville"] = ville; break
    if not criteres["ville"]:
        m = re.search(r'(?:à|a|de|depuis|en)\s+([A-ZÀ-Ÿa-zà-ÿ]{3,})', question_orig)
        if m:
            v = m.group(1).lower().strip()
            if v not in STOP_WORDS and len(v) > 3: criteres["ville"] = v

    m = re.search(
        r'(?:cherche|recherche|profil\s+de|poste\s+de|cv\s+de)\s+'
        r'(?:un|une)?\s*([a-zà-ÿ\s\-]{3,30?})'
        r'(?:\s+(?:avec|qui|de|à|pour|et|,|$))', q
    )
    if m:
        poste_brut = m.group(1).strip()
        if len(poste_brut) > 2 and poste_brut not in STOP_WORDS:
            criteres["poste"] = poste_brut
    if not criteres["poste"]:
        for metier in METIERS:
            if re.search(r'\b' + re.escape(metier) + r'\b', q, re.IGNORECASE):
                criteres["poste"] = metier; break

    for tech in TECHS:
        if re.search(r'\b' + re.escape(tech) + r'\b', q, re.IGNORECASE):
            criteres["competences"].append(tech)

    mots_caps = [
        w for w in re.findall(r'\b[A-ZÀ-Ÿ][a-zà-ÿ]{1,}\b', question_orig)
        if w.lower() not in {"Cherche","Recherche","Bonjour","CV","Profil",
                              "Candidat","Master","BTS","Abidjan","Bouaké"}
        and w.lower() not in STOP_WORDS
    ]
    if len(mots_caps) >= 2:
        criteres["prenom_nom"] = (mots_caps[0], mots_caps[1])

    mots_libres_set = set()
    for mot in re.findall(r'\b[a-zà-ÿ]{3,}\b', q):
        if mot not in STOP_WORDS and mot not in {"cherche","recherche","profil","candidat","cv","qui","aime"} and len(mot) >= 3:
            mots_libres_set.add(mot)
    criteres["mots_cles_libres"] = list(mots_libres_set)

    return criteres


# ================================================================
#  CONSTRUCTION SQL v7.1
# ================================================================

def construire_sql(criteres: dict) -> str:
    def s(v): return v.replace("'", "\\'") if v else ""

    conds = ["c.statut = 'actif'"]

    if criteres.get("genre"):
        conds.append(f"LOWER(c.genre) = '{s(criteres['genre'].lower())}'")
    if criteres.get("exp_min") is not None:
        conds.append(f"c.experience_ans >= {int(criteres['exp_min'])}")
    if criteres.get("exp_max") is not None:
        conds.append(f"c.experience_ans <= {int(criteres['exp_max'])}")
    if criteres.get("disponibilite"):
        conds.append(f"c.disponibilite = '{s(criteres['disponibilite'])}'")
    if criteres.get("niveau"):
        conds.append(f"LOWER(c.niveau_etude) LIKE LOWER('%{s(criteres['niveau'])}%')")

    if criteres.get("ville"):
        v = s(criteres["ville"].lower())
        conds.append(f"(LOWER(c.ville) LIKE '%{v}%' OR LOWER(cv.resume_ia) LIKE '%{v}%' OR LOWER(cv.texte_complet) LIKE '%{v}%')")

    if criteres.get("poste"):
        p = s(criteres["poste"].lower())
        conds.append(f"(LOWER(c.poste_actuel) LIKE '%{p}%' OR LOWER(c.competences) LIKE '%{p}%' OR LOWER(c.bio) LIKE '%{p}%' OR LOWER(cv.resume_ia) LIKE '%{p}%' OR LOWER(cv.texte_complet) LIKE '%{p}%')")

    for comp in (criteres.get("competences") or []):
        if comp and comp.strip():
            c = s(comp.strip().lower())
            conds.append(f"(LOWER(c.competences) LIKE '%{c}%' OR LOWER(c.bio) LIKE '%{c}%' OR LOWER(cv.resume_ia) LIKE '%{c}%' OR LOWER(cv.texte_complet) LIKE '%{c}%')")

    if criteres.get("prenom_nom"):
        prenom, nom = criteres["prenom_nom"]
        p_safe = s(prenom.lower()); n_safe = s(nom.lower())
        p_na = s(sans_accent(p_safe)); n_na = s(sans_accent(n_safe))
        conds.append(
            f"((LOWER(c.prenom) LIKE '%{p_safe}%' AND LOWER(c.nom) LIKE '%{n_safe}%') "
            f"OR (LOWER(c.prenom) LIKE '%{p_na}%' AND LOWER(c.nom) LIKE '%{n_na}%') "
            f"OR (LOWER(c.prenom) LIKE '%{n_safe}%' AND LOWER(c.nom) LIKE '%{p_safe}%') "
            f"OR (LOWER(c.prenom) LIKE '%{n_na}%' AND LOWER(c.nom) LIKE '%{p_na}%') "
            f"OR LOWER(cv.texte_complet) LIKE '%{p_safe}%' OR LOWER(cv.texte_complet) LIKE '%{p_na}%' "
            f"OR LOWER(cv.texte_complet) LIKE '%{n_safe}%' OR LOWER(cv.texte_complet) LIKE '%{n_na}%')"
        )

    for mot in (criteres.get("mots_cles_libres") or []):
        if mot and len(mot.strip()) >= 3:
            m = s(mot.strip().lower())
            if not any(f"'%{m}%'" in c for c in conds):
                conds.append(f"(LOWER(c.poste_actuel) LIKE '%{m}%' OR LOWER(c.competences) LIKE '%{m}%' OR LOWER(c.bio) LIKE '%{m}%' OR LOWER(c.ville) LIKE '%{m}%' OR LOWER(cv.resume_ia) LIKE '%{m}%' OR LOWER(cv.texte_complet) LIKE '%{m}%')")

    where = " AND ".join(conds)
    return (
        "SELECT c.id, c.prenom, c.nom, c.email, c.telephone, c.ville, "
        "c.genre, c.poste_actuel, c.experience_ans, c.competences, "
        "c.disponibilite, c.niveau_etude, c.photo, c.date_naissance, c.bio, "
        "cv.chemin AS cv_chemin, cv.nom_fichier AS cv_nom, "
        "cv.type_fichier, cv.resume_ia, cv.statut_analyse, "
        "LEFT(cv.texte_complet, 500) AS extrait_texte "
        "FROM candidats c "
        "LEFT JOIN cv ON cv.id = (SELECT id FROM cv WHERE candidat_id = c.id ORDER BY date_upload DESC LIMIT 1) "
        f"WHERE {where} ORDER BY c.experience_ans DESC LIMIT 50"
    )


def construire_sql_elargi(criteres: dict) -> str:
    def s(v): return v.replace("'", "\\'") if v else ""

    conds = ["c.statut = 'actif'"]

    if criteres.get("poste"):
        p = s(criteres["poste"].lower())
        conds.append(f"(LOWER(c.poste_actuel) LIKE '%{p}%' OR LOWER(c.competences) LIKE '%{p}%' OR LOWER(cv.resume_ia) LIKE '%{p}%' OR LOWER(cv.texte_complet) LIKE '%{p}%')")

    if criteres.get("prenom_nom"):
        prenom, nom = criteres["prenom_nom"]
        p_safe = s(prenom.lower()); n_safe = s(nom.lower())
        p_na = s(sans_accent(p_safe)); n_na = s(sans_accent(n_safe))
        conds.append(f"(LOWER(c.prenom) LIKE '%{p_safe}%' OR LOWER(c.prenom) LIKE '%{p_na}%' OR LOWER(c.nom) LIKE '%{n_safe}%' OR LOWER(c.nom) LIKE '%{n_na}%')")

    for mot in sorted([m for m in (criteres.get("mots_cles_libres") or []) if len(m) > 3], key=len, reverse=True)[:4]:
        m = s(mot.lower())
        conds.append(f"(LOWER(c.poste_actuel) LIKE '%{m}%' OR LOWER(c.competences) LIKE '%{m}%' OR LOWER(c.bio) LIKE '%{m}%' OR LOWER(cv.resume_ia) LIKE '%{m}%' OR LOWER(cv.texte_complet) LIKE '%{m}%')")

    where = " AND ".join(conds) if len(conds) > 1 else "c.statut = 'actif'"
    return (
        "SELECT c.id, c.prenom, c.nom, c.email, c.telephone, c.ville, "
        "c.genre, c.poste_actuel, c.experience_ans, c.competences, "
        "c.disponibilite, c.niveau_etude, c.photo, c.date_naissance, c.bio, "
        "cv.chemin AS cv_chemin, cv.nom_fichier AS cv_nom, "
        "cv.type_fichier, cv.resume_ia, cv.statut_analyse, "
        "LEFT(cv.texte_complet, 500) AS extrait_texte "
        "FROM candidats c "
        "LEFT JOIN cv ON cv.id = (SELECT id FROM cv WHERE candidat_id = c.id ORDER BY date_upload DESC LIMIT 1) "
        f"WHERE {where} ORDER BY c.experience_ans DESC LIMIT 20"
    )


# ================================================================
#  BASE DE DONNÉES
# ================================================================

def get_db():
    return mysql.connector.connect(**DB)

def propre(row: dict) -> dict:
    out = {}
    for k, v in row.items():
        if isinstance(v, datetime): out[k] = v.isoformat()
        elif hasattr(v, "__float__"): out[k] = float(v)
        else: out[k] = v
    return out


# ================================================================
#  EXTRACTION TEXTE — v7.2 CLOUDINARY + LOCAL
#
#  STRATÉGIE :
#    1. URL http(s) → téléchargement dans /tmp, extraction, suppression
#    2. Chemin local → extraction directe (legacy XAMPP / Render volume)
#
#  Le fichier temporaire est TOUJOURS supprimé dans le bloc finally,
#  même si l'extraction échoue.
# ================================================================

def telecharger_depuis_cloudinary(url: str) -> tuple:
    """
    Télécharge le fichier depuis Cloudinary et le sauvegarde dans /tmp.
    Retourne (chemin_tmp, extension).
    """
    # Déterminer l'extension depuis l'URL (avant le ? éventuel)
    url_propre = url.split("?")[0].lower()
    ext = ".pdf"  # défaut — la majorité des CV sont des PDF
    for candidate_ext in [".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png"]:
        if candidate_ext in url_propre:
            ext = candidate_ext
            break

    log.info(f"Téléchargement Cloudinary : {url[:80]}…")
    resp = http_requests.get(url, timeout=60)
    resp.raise_for_status()

    # Affiner l'extension via Content-Type si besoin
    ct = resp.headers.get("Content-Type", "")
    if "pdf" in ct:
        ext = ".pdf"
    elif "officedocument.wordprocessingml" in ct or "msword" in ct:
        ext = ".docx"
    elif "jpeg" in ct or "jpg" in ct:
        ext = ".jpg"
    elif "png" in ct:
        ext = ".png"

    fd, chemin_tmp = tempfile.mkstemp(suffix=ext, dir="/tmp")
    with os.fdopen(fd, "wb") as f:
        f.write(resp.content)

    log.info(f"Fichier temporaire : {chemin_tmp} ({len(resp.content)//1024} Ko, ext={ext})")
    return chemin_tmp, ext


def extraire_texte_pdf(chemin: str) -> str:
    texte = ""
    try:
        import pdfplumber
        with pdfplumber.open(chemin) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: texte += t + "\n"
    except Exception as e:
        log.warning(f"pdfplumber: {e}")
    # Fallback OCR si PDF scanné
    if len(texte.strip()) < 50:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            if TESSERACT_CMD: pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
            for img in convert_from_path(chemin, dpi=200):
                texte += pytesseract.image_to_string(img, lang="fra+eng") + "\n"
        except Exception as e:
            log.warning(f"OCR PDF: {e}")
    return texte.strip()


def extraire_texte_docx(chemin: str) -> str:
    texte = ""
    try:
        from docx import Document
        doc = Document(chemin)
        for p in doc.paragraphs: texte += p.text + "\n"
        for t in doc.tables:
            for r in t.rows: texte += " ".join(c.text for c in r.cells) + "\n"
    except Exception as e:
        log.warning(f"docx: {e}")
    return texte.strip()


def extraire_texte_image(chemin: str) -> str:
    try:
        import pytesseract; from PIL import Image
        if TESSERACT_CMD: pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
        return pytesseract.image_to_string(Image.open(chemin), lang="fra+eng").strip()
    except Exception as e:
        log.warning(f"OCR image: {e}"); return ""


def extraire_texte_depuis_chemin(chemin: str) -> str:
    """Dispatch selon l'extension du fichier local."""
    ext = Path(chemin).suffix.lower()
    if ext == ".pdf":            return extraire_texte_pdf(chemin)
    if ext in (".docx", ".doc"): return extraire_texte_docx(chemin)
    if ext in (".jpg", ".jpeg", ".png"): return extraire_texte_image(chemin)
    log.warning(f"Extension non supportée : {ext}")
    return ""


def extraire_texte_fichier(chemin_ou_url: str) -> str:
    """
    Point d'entrée unique — URL Cloudinary OU chemin local.
    v7.2 : les deux cas sont gérés ici, extraire_texte_local() supprimée.
    """
    if not chemin_ou_url:
        return ""

    # ── CAS 1 : URL Cloudinary (http / https) ────────────────────
    if chemin_ou_url.startswith("http://") or chemin_ou_url.startswith("https://"):
        chemin_tmp = None
        try:
            chemin_tmp, _ = telecharger_depuis_cloudinary(chemin_ou_url)
            texte = extraire_texte_depuis_chemin(chemin_tmp)
            if not texte:
                log.warning(f"Texte vide après extraction Cloudinary : {chemin_ou_url[:60]}")
            return texte
        except Exception as e:
            log.error(f"Erreur extraction Cloudinary : {e}")
            return ""
        finally:
            # Suppression garantie du fichier temporaire
            if chemin_tmp and os.path.exists(chemin_tmp):
                try: os.remove(chemin_tmp)
                except Exception: pass

    # ── CAS 2 : Chemin local (legacy XAMPP / Render volume) ──────
    chemin_abs = chemin_ou_url
    if RACINE_PROJET and not os.path.isabs(chemin_ou_url):
        chemin_abs = os.path.join(RACINE_PROJET, chemin_ou_url.lstrip("/\\"))

    if not os.path.exists(chemin_abs):
        log.error(f"Fichier local introuvable : {chemin_abs}")
        return ""

    return extraire_texte_depuis_chemin(chemin_abs)


# ================================================================
#  GROQ — ANALYSE CV
# ================================================================

PROMPT_ANALYSE_CV = """Analyse ce CV. Retourne UNIQUEMENT ce JSON :
{
  "poste_principal": "intitulé",
  "experience_ans": 0,
  "competences": ["c1","c2"],
  "villes_mentions": ["ville1"],
  "formations": ["diplome - ecole"],
  "langues": ["Français"],
  "secteurs": ["secteur"],
  "resume_enrichi": "3-5 phrases sur compétences, expériences, villes, formations, traits personnels"
}"""

def groq_analyser_cv(texte: str, info: dict) -> dict:
    prompt = (
        f"{PROMPT_ANALYSE_CV}\n\n"
        f"Candidat: {info.get('prenom','')} {info.get('nom','')}\n"
        f"CV:\n---\n{texte[:5000]}\n---"
    )
    contenu = groq_appel(
        messages=[{"role": "user", "content": prompt}],
        max_tokens=700, temperature=0.1, json_mode=True,
    )
    try:
        return json.loads(nettoyer_json(contenu))
    except Exception as e:
        log.error(f"Analyse CV JSON: {e}"); return {}


def fusionner_competences(ex: str, nv: str) -> str:
    connus = {c.strip().lower() for c in ex.split(",") if c.strip()}
    res    = [c for c in ex.split(",") if c.strip()]
    for c in nv.split(","):
        if c.strip() and c.strip().lower() not in connus:
            res.append(c.strip()); connus.add(c.strip().lower())
    return ", ".join(res)


def sauvegarder_analyse(cur, cv_id: int, candidat_id: int,
                         donnees: dict, texte_complet: str,
                         comp_ex: str, bio_ex: str):
    resume = donnees.get("resume_enrichi", "")
    comp   = fusionner_competences(comp_ex or "", ", ".join(donnees.get("competences", [])))
    resume_ia = (
        f"POSTE: {donnees.get('poste_principal','')}\n"
        f"SECTEURS: {', '.join(donnees.get('secteurs',[]))}\n"
        f"VILLES: {', '.join(donnees.get('villes_mentions',[]))}\n"
        f"FORMATIONS: {', '.join(donnees.get('formations',[]))}\n"
        f"LANGUES: {', '.join(donnees.get('langues',[]))}\n"
        f"RESUME: {resume}"
    )
    texte_a_stocker = (texte_complet or "")[:100_000]
    cur.execute(
        "UPDATE cv SET statut_analyse='analyse', resume_ia=%s, texte_complet=%s WHERE id=%s",
        (resume_ia, texte_a_stocker, cv_id)
    )
    cur.execute("""
        UPDATE candidats SET competences=%s,
            bio=COALESCE(NULLIF(bio,''),%s),
            poste_actuel=COALESCE(NULLIF(poste_actuel,''),%s),
            experience_ans=CASE WHEN experience_ans IS NULL OR experience_ans=0
                           THEN %s ELSE experience_ans END
        WHERE id=%s
    """, (comp, resume, donnees.get("poste_principal",""),
          int(donnees.get("experience_ans", 0)), candidat_id))


def analyser_cv_batch() -> int:
    """
    Analyse les CV en attente.
    v7.2 : cv.chemin peut être une URL Cloudinary → géré nativement.
    """
    db  = get_db(); cur = db.cursor(dictionary=True)
    cur.execute("""
        SELECT cv.id AS cv_id, cv.chemin, cv.candidat_id,
               c.prenom, c.nom, c.poste_actuel, c.ville, c.competences, c.bio
        FROM cv JOIN candidats c ON c.id = cv.candidat_id
        WHERE cv.statut_analyse = 'en_attente'
          AND cv.chemin IS NOT NULL AND cv.chemin != ''
        LIMIT 10
    """)
    cvs = cur.fetchall(); nb = 0
    for cv in cvs:
        log.info(f"Batch CV #{cv['cv_id']} — {cv['chemin'][:60]}")
        texte = extraire_texte_fichier(cv["chemin"])
        if not texte:
            log.warning(f"CV #{cv['cv_id']} : texte vide → erreur")
            cur.execute("UPDATE cv SET statut_analyse='erreur' WHERE id=%s", (cv["cv_id"],))
            db.commit(); continue
        try:
            donnees = groq_analyser_cv(texte, cv)
        except HTTPException:
            break
        if donnees:
            sauvegarder_analyse(cur, cv["cv_id"], cv["candidat_id"], donnees, texte,
                                cv.get("competences") or "", cv.get("bio") or "")
            db.commit(); nb += 1
            log.info(f"✅ CV #{cv['cv_id']} analysé — {donnees.get('poste_principal','?')}")
    cur.close(); db.close()
    return nb


# ================================================================
#  GROQ — SCORING
# ================================================================

SYSTEM_SCORING = """Tu es expert RH. Évalue la pertinence d'un profil pour une recherche.
Retourne UNIQUEMENT ce JSON :
{
  "score": <0-100>,
  "justification": "<1 phrase>",
  "points_forts": "<points forts>",
  "points_faibles": "<points faibles>",
  "competences_detectees": "<compétences clés>"
}
Barème : 90+=parfait | 70-89=bon | 50-69=correct | 30-49=faible | <30=hors sujet
- Si cv_analyse rempli : s'y fier en priorité sur le formulaire
- Synonymes : JS=JavaScript, dev=développeur, comm=communication
- Profil vide → score 0-15""".strip()

def groq_scorer(candidat: dict, question: str, criteres: dict) -> tuple:
    cv_info     = (candidat.get("resume_ia") or "")[:300]
    extrait     = (candidat.get("extrait_texte") or "")[:200]
    cv_contexte = cv_info + ("\n[Extrait CV brut]: " + extrait if extrait else "")
    profil = {
        "poste":       candidat.get("poste_actuel", "—"),
        "exp":         f"{candidat.get('experience_ans', 0)} ans",
        "ville":       candidat.get("ville", "—"),
        "competences": (candidat.get("competences") or "")[:150],
        "cv":          cv_contexte or "non analysé",
    }
    msg = (f'RECHERCHE: "{question}"\nFILTRE: {criteres.get("ville") or "aucun"}\n'
           f"PROFIL: {json.dumps(profil, ensure_ascii=False)}")
    try:
        contenu = groq_appel(
            messages=[{"role":"system","content":SYSTEM_SCORING},{"role":"user","content":msg}],
            max_tokens=180, temperature=0.1, json_mode=True,
        )
        r     = json.loads(nettoyer_json(contenu))
        score = max(0, min(100, int(r.get("score", 0))))
        res   = r.get("justification", "")
        if r.get("points_forts"):          res += f" | ✅ {r['points_forts']}"
        if r.get("competences_detectees"): res += f" | 🔧 {r['competences_detectees']}"
        return score, res[:350]
    except Exception as e:
        log.warning(f"Scorer: {e}"); return 50, "Profil potentiellement pertinent."


# ================================================================
#  GROQ — RÉPONSE NATURELLE
# ================================================================

def groq_reponse(question: str, nb: int, elargi: bool = False) -> str:
    if nb > 0:
        extra = " (critères assouplis)" if elargi else ""
        return f"J'ai trouvé {nb} candidat{'s' if nb>1 else ''} correspondant{extra}."
    try:
        db  = get_db(); cur = db.cursor(dictionary=True)
        cur.execute("SELECT c.prenom, c.nom, c.poste_actuel, c.ville, c.experience_ans "
                    "FROM candidats c WHERE c.statut='actif' ORDER BY c.date_inscription DESC LIMIT 5")
        profils = cur.fetchall(); cur.close(); db.close()
        if not profils: return "Aucun candidat dans la base."
        liste = "\n".join(f"- {p['prenom']} {p['nom']} | {p['poste_actuel'] or '?'} | {p['ville'] or '?'}" for p in profils)
        return groq_appel(
            messages=[{"role":"user","content":(
                f'Recherche: "{question}"\nAucun résultat.\n\nProfils disponibles:\n{liste}\n\n'
                "En 3 phrases: explique l'absence, propose 1-2 profils proches, suggère comment reformuler."
            )}], max_tokens=200, temperature=0.3,
        ).strip()
    except Exception:
        return f"Aucun candidat ne correspond à « {question} »."


# ================================================================
#  CV DE PROPOSITION — v7.1
# ================================================================

def recuperer_cv_proposition(limite: int = 5) -> list:
    try:
        db  = get_db(); cur = db.cursor(dictionary=True)
        cur.execute("""
            SELECT c.id, c.prenom, c.nom, c.email, c.telephone, c.ville,
                   c.genre, c.poste_actuel, c.experience_ans, c.competences,
                   c.disponibilite, c.niveau_etude, c.photo, c.date_naissance, c.bio,
                   cv.chemin AS cv_chemin, cv.nom_fichier AS cv_nom,
                   cv.type_fichier, cv.resume_ia, cv.statut_analyse
            FROM candidats c
            LEFT JOIN cv ON cv.id = (SELECT id FROM cv WHERE candidat_id = c.id ORDER BY date_upload DESC LIMIT 1)
            WHERE c.statut = 'actif' ORDER BY c.date_inscription DESC LIMIT %s
        """, (limite,))
        rows = cur.fetchall(); cur.close(); db.close()
        propositions = []
        for c in rows:
            cp = propre(dict(c))
            propositions.append({**cp, "score": 0, "resume_ia": cp.get("resume_ia",""),
                                  "nom": f"{cp.get('prenom','')} {cp.get('nom','')}".strip(),
                                  "profession": cp.get("poste_actuel",""),
                                  "experience": f"{cp.get('experience_ans',0)} ans",
                                  "proposition": True})
        return propositions
    except Exception as e:
        log.warning(f"Proposition fallback: {e}"); return []


# ================================================================
#  STATISTIQUES
# ================================================================

def repondre_statistiques(question: str) -> str:
    db  = get_db(); cur = db.cursor(dictionary=True); stats = {}
    try:
        cur.execute("SELECT COUNT(*) AS n FROM candidats WHERE statut='actif'")
        stats["candidats"] = cur.fetchone()["n"]
        cur.execute("SELECT COUNT(*) AS n FROM cv WHERE statut_analyse='analyse'")
        stats["cv_analyses"] = cur.fetchone()["n"]
        cur.execute("SELECT COUNT(*) AS n FROM cv WHERE statut_analyse='en_attente'")
        stats["cv_en_attente"] = cur.fetchone()["n"]
        cur.execute("SELECT COUNT(*) AS n FROM candidats WHERE LOWER(genre)='femme' AND statut='actif'")
        stats["femmes"] = cur.fetchone()["n"]
        cur.execute("SELECT COUNT(*) AS n FROM candidats WHERE LOWER(genre)='homme' AND statut='actif'")
        stats["hommes"] = cur.fetchone()["n"]
        cur.execute("SELECT AVG(experience_ans) AS m FROM candidats WHERE statut='actif'")
        moy = cur.fetchone()["m"]
        stats["exp_moy"] = round(float(moy), 1) if moy else 0
        cur.execute("SELECT ville, COUNT(*) AS n FROM candidats WHERE statut='actif' GROUP BY ville ORDER BY n DESC LIMIT 3")
        stats["top_villes"] = [f"{r['ville']}({r['n']})" for r in cur.fetchall() if r["ville"]]
    finally:
        cur.close(); db.close()
    return groq_appel(
        messages=[{"role":"user","content":f"Stats: {json.dumps(stats, ensure_ascii=False)}\nQuestion: {question}\nRéponds en français, 2-3 phrases."}],
        max_tokens=150, temperature=0.2,
    ).strip()


# ================================================================
#  ENDPOINT PRINCIPAL — /ask
# ================================================================

@app.post("/ask", tags=["IA"])
@app.post("/agent", tags=["IA"])
async def ask(body: RequeteChat):
    question = body.question.strip()
    if not question: raise HTTPException(400, "Question vide.")

    log.info(f"/ask — '{question[:70]}'")
    question_norm = normaliser(question)
    intention     = classifier(question_norm)
    log.info(f"Intention: {intention} | Normalisée: {question_norm[:60]}")

    if intention == "conversation":
        return {"success": True, "type": "conversation",
                "reponse": "Bonjour ! Décrivez le profil recherché. Exemple : développeur PHP 3 ans Abidjan.",
                "results": [], "propositions": [], "sql_generated": ""}

    if intention == "statistique":
        return {"success": True, "type": "statistique",
                "reponse": repondre_statistiques(question),
                "results": [], "propositions": [], "sql_generated": ""}

    criteres = extraire_criteres(question_norm, question)
    log.info(f"Critères: poste={criteres.get('poste')} | ville={criteres.get('ville')} | "
             f"genre={criteres.get('genre')} | exp_min={criteres.get('exp_min')} | "
             f"prenom_nom={criteres.get('prenom_nom')} | mots_cles={criteres.get('mots_cles_libres')}")

    sql = construire_sql(criteres); elargi = False

    try:
        db  = get_db(); cur = db.cursor(dictionary=True)
        cur.execute(sql); candidats = cur.fetchall()
        cur.close(); db.close()
    except mysql.connector.Error as e:
        log.error(f"SQL erreur: {e}"); raise HTTPException(500, f"Erreur base de données: {e}")

    log.info(f"{len(candidats)} candidat(s) — SQL strict")

    if not candidats and (criteres.get("poste") or criteres.get("prenom_nom") or criteres.get("mots_cles_libres")):
        sql_e = construire_sql_elargi(criteres)
        try:
            db  = get_db(); cur = db.cursor(dictionary=True)
            cur.execute(sql_e); candidats = cur.fetchall(); cur.close(); db.close()
            if candidats: elargi = True; sql = sql_e; log.info(f"{len(candidats)} candidat(s) — SQL élargi")
        except mysql.connector.Error as e:
            log.warning(f"SQL élargi erreur: {e}")

    resultats = []
    for cand in candidats:
        c = propre(dict(cand))
        score, resume = groq_scorer(c, question_norm, criteres)
        if score >= SCORE_MIN:
            resultats.append({**c, "score": score, "resume_ia": resume,
                               "nom": f"{c.get('prenom','')} {c.get('nom','')}".strip(),
                               "profession": c.get("poste_actuel",""),
                               "experience": f"{c.get('experience_ans',0)} ans"})

    resultats.sort(key=lambda x: x["score"], reverse=True)
    nb = len(resultats)

    candidats_proposition = []
    if nb == 0:
        candidats_proposition = recuperer_cv_proposition(limite=5)
        log.info(f"{len(candidats_proposition)} CV de proposition récupérés")

    reponse = groq_reponse(question_norm, nb, elargi)
    return {"success": True, "type": "recherche", "reponse": reponse,
            "results": resultats, "propositions": candidats_proposition,
            "count": nb, "sql_generated": sql}


# ================================================================
#  ENDPOINTS CV
# ================================================================

@app.post("/cv-uploade", tags=["CV"])
async def cv_uploade(body: RequeteUploadCV):
    """
    Analyse automatique après upload PHP.
    v7.2 : priorité cloudinary_url > chemin_cv > texte_cv.
    """
    db  = get_db(); cur = db.cursor(dictionary=True)
    cur.execute("SELECT * FROM candidats WHERE id=%s", (body.candidat_id,))
    candidat = cur.fetchone()
    if not candidat:
        cur.close(); db.close(); raise HTTPException(404, f"Candidat #{body.candidat_id} introuvable.")

    cur.execute("SELECT id FROM cv WHERE candidat_id=%s ORDER BY date_upload DESC LIMIT 1", (body.candidat_id,))
    cv_row = cur.fetchone()
    if not cv_row:
        cur.close(); db.close(); raise HTTPException(404, "Aucun CV en base.")

    # Priorité : URL Cloudinary > chemin local > texte pré-extrait
    texte = ""
    source = (body.cloudinary_url or body.chemin_cv or "").strip()
    if source:
        texte = extraire_texte_fichier(source)
    if not texte:
        texte = (body.texte_cv or "").strip()

    if not texte:
        cur.execute("UPDATE cv SET statut_analyse='erreur' WHERE id=%s", (cv_row["id"],))
        db.commit(); cur.close(); db.close()
        return {"success": False, "message": "Fichier illisible ou texte vide."}

    donnees = groq_analyser_cv(texte, dict(candidat))
    if not donnees:
        cur.execute("UPDATE cv SET statut_analyse='erreur' WHERE id=%s", (cv_row["id"],))
        db.commit(); cur.close(); db.close()
        return {"success": False, "message": "Analyse Groq échouée."}

    sauvegarder_analyse(cur, cv_row["id"], body.candidat_id, donnees, texte,
                        candidat.get("competences") or "", candidat.get("bio") or "")
    db.commit(); cur.close(); db.close()
    log.info(f"✅ CV analysé #{body.candidat_id}: {donnees.get('poste_principal','?')}")
    return {"success": True, "message": f"CV analysé : {donnees.get('poste_principal','?')}",
            "competences": ", ".join(donnees.get("competences", [])),
            "resume": donnees.get("resume_enrichi", "")}


@app.post("/analyser-cv-fichier", tags=["CV"])
async def analyser_cv_fichier(body: RequeteAnalyserFichier):
    """
    Analyse un CV — chemin_cv peut être une URL Cloudinary complète.
    v7.2 : extraire_texte_fichier() gère les deux cas.
    """
    db  = get_db(); cur = db.cursor(dictionary=True)
    cur.execute("SELECT * FROM candidats WHERE id=%s", (body.candidat_id,))
    candidat = cur.fetchone()
    if not candidat:
        cur.close(); db.close(); raise HTTPException(404, f"Candidat #{body.candidat_id} introuvable.")

    cur.execute("SELECT id FROM cv WHERE candidat_id=%s ORDER BY date_upload DESC LIMIT 1", (body.candidat_id,))
    cv_row = cur.fetchone()
    if not cv_row:
        cur.close(); db.close(); raise HTTPException(404, "Aucun CV trouvé.")

    texte = extraire_texte_fichier(body.chemin_cv)
    if not texte:
        cur.execute("UPDATE cv SET statut_analyse='erreur' WHERE id=%s", (cv_row["id"],))
        db.commit(); cur.close(); db.close()
        return {"success": False, "message": f"Fichier illisible : {body.chemin_cv[:60]}"}

    donnees = groq_analyser_cv(texte, dict(candidat))
    if not donnees:
        cur.execute("UPDATE cv SET statut_analyse='erreur' WHERE id=%s", (cv_row["id"],))
        db.commit(); cur.close(); db.close()
        return {"success": False, "message": "Analyse Groq échouée."}

    sauvegarder_analyse(cur, cv_row["id"], body.candidat_id, donnees, texte,
                        candidat.get("competences") or "", candidat.get("bio") or "")
    db.commit(); cur.close(); db.close()
    log.info(f"✅ CV analysé #{body.candidat_id}: {donnees.get('poste_principal','?')}")
    return {"success": True, "message": f"CV analysé : {donnees.get('poste_principal','?')}",
            "competences": ", ".join(donnees.get("competences", [])),
            "resume": donnees.get("resume_enrichi", "")}


@app.post("/analyser-cv-en-attente", tags=["CV"])
async def analyser_cv_en_attente():
    nb = analyser_cv_batch()
    return {"success": True, "traites": nb,
            "message": f"{nb} CV analysé(s)." if nb else "Aucun CV en attente."}


# ================================================================
#  UTILITAIRES
# ================================================================

@app.get("/cvs", tags=["Données"])
async def get_all_cvs(limit: int = 100):
    try:
        db  = get_db(); cur = db.cursor(dictionary=True)
        cur.execute("""
            SELECT c.id, c.prenom, c.nom, c.email, c.ville, c.poste_actuel,
                   c.experience_ans, c.competences, c.photo,
                   cv.chemin AS cv_chemin, cv.nom_fichier AS cv_nom,
                   cv.resume_ia, cv.statut_analyse,
                   (cv.texte_complet IS NOT NULL AND cv.texte_complet != '') AS a_texte_complet
            FROM candidats c
            LEFT JOIN cv ON cv.id=(SELECT id FROM cv WHERE candidat_id=c.id ORDER BY date_upload DESC LIMIT 1)
            WHERE c.statut='actif' ORDER BY c.date_inscription DESC LIMIT %s
        """, (limit,))
        rows = [propre(r) for r in cur.fetchall()]; cur.close(); db.close()
        return {"success": True, "count": len(rows), "cvs": rows}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/sante", tags=["Système"])
async def sante():
    bd_ok = False; cv_attente = 0; cv_avec_texte = 0
    try:
        db  = get_db(); cur = db.cursor()
        cur.execute("SELECT COUNT(*) FROM cv WHERE statut_analyse='en_attente'")
        cv_attente = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM cv WHERE texte_complet IS NOT NULL AND texte_complet != ''")
        cv_avec_texte = cur.fetchone()[0]
        cur.close(); db.close(); bd_ok = True
    except Exception: pass
    return {"status": "ok" if bd_ok else "erreur_bd", "version": "7.2",
            "bd": bd_ok, "groq": bool(GROQ_API_KEY),
            "cv_en_attente": cv_attente, "cv_avec_texte": cv_avec_texte,
            "cloudinary": "✅ URLs http/https supportées nativement",
            "render": "✅ /tmp utilisé pour les fichiers temporaires"}


@app.get("/", tags=["Système"])
async def root():
    return {
        "service": "SmartRecruit Matching IA v7.2",
        "nouveautes_v72": [
            "Cloudinary : téléchargement automatique depuis URL http(s) dans /tmp",
            "extraire_texte_fichier() unifié : URL Cloudinary OU chemin local",
            "Bug corrigé : extraire_texte_local() supprimée (n'existait pas)",
            "RequeteUploadCV : nouveau champ cloudinary_url",
            "Timeout 60s pour téléchargement Cloudinary",
            "Nettoyage automatique du fichier /tmp après extraction",
        ],
        "endpoints": ["/ask", "/cv-uploade", "/analyser-cv-fichier",
                      "/analyser-cv-en-attente", "/cvs", "/sante", "/docs"],
    }


# ================================================================
#  DÉMARRAGE
# ================================================================

if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*60)
    print("  CvMatchIA Matching IA v7.2")
    print(f"  Base      : {DB['database']}")
    print(f"  Groq      : {'✅' if GROQ_API_KEY else '❌ clé manquante'}")
    print(f"  Cloudinary: ✅ URLs http/https supportées nativement")
    print(f"  Render    : ✅ /tmp pour les fichiers temporaires")
    print("  Docs      : http://localhost:5001/docs")
    print("="*60)
    print("\n  Analyse des CV en attente...")
    nb = analyser_cv_batch()
    print(f"  {'✅ '+str(nb)+' CV(s) analysé(s)' if nb else 'ℹ️  Aucun CV en attente'}\n")
    uvicorn.run("app.py:app", host="0.0.0.0", port=5001, reload=True)
